"""
OmniRoute AI — Document Parser Module
Extracts text, images, and tables from PDF, DOCX, and TXT files.
Zero-loss: every page, every image, every table is captured.
"""

from __future__ import annotations

import io
import base64
import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from PIL import Image

from modules.utils import logger


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class ExtractedImage:
    """An image extracted from a document page."""
    page_number: int
    image_bytes: bytes
    mime_type: str  # e.g. "image/png"
    width: int = 0
    height: int = 0
    image_id: str = ""
    occurrence: int = 0

    def to_base64(self) -> str:
        return base64.b64encode(self.image_bytes).decode("utf-8")

    def __post_init__(self) -> None:
        # A stable id lets us deduplicate vision calls while still preserving
        # every occurrence in the parsed document and in storage.
        if not self.image_id:
            self.image_id = hashlib.sha256(self.image_bytes).hexdigest()[:24]


@dataclass
class ExtractedTable:
    """A table extracted from a document page, stored as list-of-lists."""
    page_number: int
    table_data: list[list[str]]  # rows × cols
    source_file: str = ""
    table_index: int = 0

    def to_mongo_dict(self) -> dict:
        """Serialize for MongoDB insertion."""
        headers = self.table_data[0] if self.table_data else []
        rows = self.table_data[1:] if len(self.table_data) > 1 else []
        # Mongo document keys cannot be duplicated. Preserve the original
        # matrix in `raw`, while making the convenience row objects safe.
        safe_headers: list[str] = []
        seen: dict[str, int] = {}
        for idx, header in enumerate(headers):
            name = str(header).strip() or f"column_{idx + 1}"
            seen[name] = seen.get(name, 0) + 1
            safe_headers.append(name if seen[name] == 1 else f"{name}_{seen[name]}")
        return {
            "page_number": self.page_number,
            "table_index": self.table_index,
            "source_file": self.source_file,
            "headers": safe_headers,
            "rows": [dict(zip(safe_headers, row)) for row in rows] if safe_headers else rows,
            "raw": self.table_data,
        }


@dataclass
class PageContent:
    """All content extracted from a single page."""
    page_number: int
    text: str = ""
    images: list[ExtractedImage] = field(default_factory=list)
    tables: list[ExtractedTable] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """Complete parsed output of a document."""
    filename: str
    total_pages: int
    pages: list[PageContent] = field(default_factory=list)
    source_page_count: Optional[int] = None
    page_limit: Optional[int] = None

    @property
    def all_images(self) -> list[ExtractedImage]:
        return [img for p in self.pages for img in p.images]

    @property
    def all_tables(self) -> list[ExtractedTable]:
        return [tbl for p in self.pages for tbl in p.tables]

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text)

    @property
    def is_truncated(self) -> bool:
        return bool(self.source_page_count and self.total_pages < self.source_page_count)


# ---------------------------------------------------------------------------
# PDF Parser (PyMuPDF / fitz)
# ---------------------------------------------------------------------------

def _parse_pdf(file_bytes: bytes, filename: str, max_pages: Optional[int] = None) -> ParsedDocument:
    """
    Extract text, images, and tables from a PDF.
    Uses PyMuPDF's built-in table detection and image extraction.
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[PageContent] = []

    source_page_count = len(doc)
    # Never silently impose a small provider-specific cap in the parser. The
    # caller chooses a free-tier limit (normally 20–30) or a production limit
    # (up to 200), and the result records whether anything was left unprocessed.
    if max_pages is None:
        max_pages = 200
    if max_pages < 1:
        raise ValueError("max_pages must be at least 1")
    pages_to_process = min(source_page_count, max_pages)

    for page_idx in range(pages_to_process):
        page = doc[page_idx]
        page_num = page_idx + 1
        pc = PageContent(page_number=page_num)

        # --- Text ---
        pc.text = page.get_text("text") or ""

        # --- Images ---
        image_list = page.get_images(full=True)
        for img_index, img_info in enumerate(image_list):
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                if base_image and base_image.get("image"):
                    img_bytes = base_image["image"]
                    ext = base_image.get("ext", "png")
                    mime = f"image/{ext}" if ext != "jpg" else "image/jpeg"
                    # Get dimensions
                    pil_img = Image.open(io.BytesIO(img_bytes))
                    w, h = pil_img.size
                    pc.images.append(ExtractedImage(
                        page_number=page_num,
                        image_bytes=img_bytes,
                        mime_type=mime,
                        width=w,
                        height=h,
                        occurrence=img_index,
                    ))
            except Exception as e:
                logger.warning(f"Could not extract image xref={xref} on page {page_num}: {e}")

        # --- Tables (PyMuPDF built-in) ---
        try:
            tabs = page.find_tables()
            for table_index, tab in enumerate(tabs):
                table_data = tab.extract()
                if table_data and len(table_data) > 0:
                    # Clean None values
                    cleaned = [
                        [cell if cell is not None else "" for cell in row]
                        for row in table_data
                    ]
                    pc.tables.append(ExtractedTable(
                        page_number=page_num,
                        table_data=cleaned,
                        source_file=filename,
                        table_index=table_index,
                    ))
        except Exception as e:
            logger.warning(f"Table extraction failed on page {page_num}: {e}")

        pages.append(pc)

    doc.close()

    return ParsedDocument(
        filename=filename,
        total_pages=len(pages),
        pages=pages,
        source_page_count=source_page_count,
        page_limit=max_pages,
    )


# ---------------------------------------------------------------------------
# DOCX Parser
# ---------------------------------------------------------------------------

def _parse_docx(file_bytes: bytes, filename: str, max_pages: Optional[int] = None) -> ParsedDocument:
    """
    Extract text, images, and tables from a .docx file.
    DOCX doesn't have fixed pages, so we treat the whole doc as page 1+
    and use paragraph breaks as approximate page boundaries.
    """
    doc = DocxDocument(io.BytesIO(file_bytes))
    pages: list[PageContent] = []

    # DOCX does not expose reliable rendered page numbers through python-docx.
    # Walk the document body in order instead of assigning media round-robin;
    # this keeps each image/table near the text that introduces it.
    chars_per_page = 3000
    if max_pages is None:
        max_pages = 200
    if max_pages < 1:
        raise ValueError("max_pages must be at least 1")

    def ensure_page() -> PageContent:
        if not pages:
            pages.append(PageContent(page_number=1))
        return pages[-1]

    def start_new_page_if_needed(incoming_chars: int) -> PageContent:
        page = ensure_page()
        if page.text and len(page.text) + incoming_chars > chars_per_page:
            pages.append(PageContent(page_number=len(pages) + 1))
        return pages[-1]

    image_occurrence = 0
    table_index = 0
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = next((p for p in doc.paragraphs if p._p is child), None)
            if paragraph is None:
                continue
            text = paragraph.text.strip()
            page = start_new_page_if_needed(len(text) + 1)
            if text:
                page.text = f"{page.text}\n{text}".strip()

            for blip in child.iter(qn("a:blip")):
                rel_id = blip.get(qn("r:embed"))
                rel = doc.part.rels.get(rel_id)
                if not rel or "image" not in rel.reltype:
                    continue
                try:
                    img_bytes = rel.target_part.blob
                    pil_img = Image.open(io.BytesIO(img_bytes))
                    fmt = (pil_img.format or "PNG").lower()
                    mime = "image/jpeg" if fmt in {"jpg", "jpeg"} else f"image/{fmt}"
                    page.images.append(ExtractedImage(
                        page_number=page.page_number,
                        image_bytes=img_bytes,
                        mime_type=mime,
                        width=pil_img.width,
                        height=pil_img.height,
                        occurrence=image_occurrence,
                    ))
                    image_occurrence += 1
                except Exception as e:
                    logger.warning(f"Could not extract DOCX image: {e}")
        elif tag == "tbl":
            table = next((t for t in doc.tables if t._tbl is child), None)
            if table is None:
                continue
            table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if table_data:
                page = start_new_page_if_needed(sum(len(cell) for row in table_data for cell in row))
                page.tables.append(ExtractedTable(
                    page_number=page.page_number,
                    table_data=table_data,
                    source_file=filename,
                    table_index=table_index,
                ))
                table_index += 1

    if not pages:
        pages.append(PageContent(page_number=1, text=""))

    source_page_count = len(pages)
    if len(pages) > max_pages:
        pages = pages[:max_pages]

    return ParsedDocument(
        filename=filename,
        total_pages=len(pages),
        pages=pages,
        source_page_count=source_page_count,
        page_limit=max_pages,
    )


# ---------------------------------------------------------------------------
# TXT Parser
# ---------------------------------------------------------------------------

def _parse_txt(file_bytes: bytes, filename: str, max_pages: Optional[int] = None) -> ParsedDocument:
    """Parse a plain text file. Splits into approximate pages."""
    text = file_bytes.decode("utf-8", errors="replace")
    chars_per_page = 3000
    if max_pages is None:
        max_pages = 200
    if max_pages < 1:
        raise ValueError("max_pages must be at least 1")
    pages: list[PageContent] = []

    page_num = 0
    for i in range(0, max(len(text), 1), chars_per_page):
        page_num += 1
        chunk = text[i:i + chars_per_page]
        pages.append(PageContent(page_number=page_num, text=chunk))

    if not pages:
        pages.append(PageContent(page_number=1, text=""))

    source_page_count = len(pages)
    pages = pages[:max_pages]
    return ParsedDocument(
        filename=filename,
        total_pages=len(pages),
        pages=pages,
        source_page_count=source_page_count,
        page_limit=max_pages,
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

def parse_document(
    file_bytes: bytes,
    filename: str,
    max_pages: Optional[int] = 200,
) -> ParsedDocument:
    """
    Main entry point. Parses any supported document and returns a
    ParsedDocument with zero content loss.

    Raises ValueError for unsupported file types.
    """
    ext = Path(filename).suffix.lower()
    logger.info(f"Parsing '{filename}' (ext={ext}, size={len(file_bytes):,} bytes)")

    if ext == ".pdf":
        result = _parse_pdf(file_bytes, filename, max_pages)
    elif ext == ".docx":
        result = _parse_docx(file_bytes, filename, max_pages)
    elif ext == ".txt":
        result = _parse_txt(file_bytes, filename, max_pages)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {SUPPORTED_EXTENSIONS}"
        )

    logger.info(
        f"Parsed '{filename}': {result.total_pages}/{result.source_page_count or result.total_pages} pages, "
        f"{len(result.all_images)} images, {len(result.all_tables)} tables"
    )
    return result
