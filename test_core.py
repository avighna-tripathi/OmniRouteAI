"""Offline regression tests for the extraction and chunking layers."""

from __future__ import annotations

import io
import asyncio
import unittest
from unittest.mock import AsyncMock, patch

import fitz
from docx import Document
from PIL import Image

from modules.chunker import chunk_document
from modules.parser import ExtractedTable, parse_document
from modules.table_store import get_table_summary_text
from modules.agents import MapOutput


class CorePipelineTests(unittest.TestCase):
    def test_text_page_limit_is_explicit_and_loss_aware(self) -> None:
        source = ("Page content. " * 700).encode("utf-8")
        parsed = parse_document(source, "sample.txt", max_pages=2)
        self.assertEqual(parsed.total_pages, 2)
        self.assertGreater(parsed.source_page_count, parsed.total_pages)
        self.assertTrue(parsed.is_truncated)

    def test_pdf_parser_keeps_all_pages_up_to_selected_limit(self) -> None:
        pdf = fitz.open()
        for idx in range(3):
            page = pdf.new_page()
            page.insert_text((72, 72), f"Page {idx + 1} marker")
        parsed = parse_document(pdf.tobytes(), "sample.pdf", max_pages=2)
        pdf.close()
        self.assertEqual(parsed.total_pages, 2)
        self.assertEqual(parsed.source_page_count, 3)
        self.assertIn("Page 2 marker", parsed.full_text)

    def test_docx_tables_and_images_follow_document_order(self) -> None:
        doc = Document()
        doc.add_paragraph("Before the table")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Metric"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Revenue"
        table.cell(1, 1).text = "42"
        doc.add_paragraph("After the table")
        image = Image.new("RGB", (12, 12), "red")
        image_bytes = io.BytesIO()
        image.save(image_bytes, format="PNG")
        doc.add_picture(io.BytesIO(image_bytes.getvalue()))
        output = io.BytesIO()
        doc.save(output)

        parsed = parse_document(output.getvalue(), "sample.docx", max_pages=2)
        self.assertEqual(len(parsed.all_tables), 1)
        self.assertEqual(parsed.all_tables[0].table_data[1][1], "42")
        self.assertEqual(len(parsed.all_images), 1)

    def test_table_context_contains_cell_values_and_chunk_validation(self) -> None:
        table = ExtractedTable(
            page_number=4,
            table_data=[["Metric", "Value"], ["Revenue", "42"]],
        )
        self.assertIn("Revenue", get_table_summary_text([table]))
        self.assertIn("42", get_table_summary_text([table]))
        with self.assertRaises(ValueError):
            chunk_document(
                parse_document(b"hello", "sample.txt"),
                [],
                chunk_size=100,
                chunk_overlap=100,
            )

    def test_pipeline_preserves_stats_without_network_calls(self) -> None:
        from modules.pipeline import run_pipeline

        async def fake_map(chunk_id, chunk_text, source_pages, model=None):
            return MapOutput(
                chunk_id=chunk_id,
                source_pages=source_pages,
                facts=[f"fact from {chunk_id}"],
                summary=chunk_text[:80],
            )

        source = ("Important paragraph with a number 42. " * 500).encode("utf-8")
        with patch("modules.pipeline.store_tables", return_value=0), \
             patch("modules.pipeline.store_images", return_value=0), \
             patch("modules.pipeline.caption_images", new=AsyncMock(return_value=[])), \
             patch("modules.pipeline.run_map_phase_single", side_effect=fake_map), \
             patch("modules.pipeline.run_executive_agent", new=AsyncMock(return_value="# Summary")), \
             patch("modules.pipeline.run_critic_agent", new=AsyncMock(return_value={
                 "is_consistent": True,
                 "quality_score": 9,
                 "feedback": "ok",
             })):
            result = asyncio.run(run_pipeline(source, "sample.txt", max_pages=1, chunk_size=1000, chunk_overlap=100))

        self.assertEqual(result.master_summary, "# Summary")
        self.assertTrue(result.is_consistent)
        self.assertTrue(result.stats["pages_truncated"])
        self.assertGreater(result.stats["total_chunks"], 0)


if __name__ == "__main__":
    unittest.main()
